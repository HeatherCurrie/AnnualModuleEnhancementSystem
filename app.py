from flask import Flask, request, jsonify, render_template, session, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import text
from flask_mysqldb import MySQL
import os
from flask_cors import CORS
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask_mail import Mail, Message
from functools import wraps
import bcrypt


load_dotenv()
app = Flask(__name__)
CORS(app)
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('SQLALCHEMY_DATABASE_URI')
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY')
SSL_CA_PATH = os.path.join(os.getcwd(), os.getenv('SSL_CA_PATH'))

app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'connect_args': {
        'ssl': {
            'ca': SSL_CA_PATH
        }
    }
}


# EMAIL CONFIG
app.config['MAIL_SERVER']='smtp.gmail.com'
app.config['MAIL_PORT'] = 465
app.config['MAIL_USERNAME'] = 'annualreviewsystem@gmail.com'
app.config['MAIL_PASSWORD'] = os.getenv('PASSWORD')
app.config['MAIL_USE_TLS'] = False
app.config['MAIL_USE_SSL'] = True
mail = Mail(app)
mysql = MySQL(app)

db = SQLAlchemy(app)


# LOGIN REQUIRED TO VIEW PAGES 
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'userID' not in session:
            # Redirect to login page if user is not logged in
            return redirect(url_for('home'))
        return f(*args, **kwargs)
    return decorated_function


# FUNCTIONS TO NAVIGATE BETWEEN TEMPLATES
# UPON STARTUP
@app.route('/')
def home():
    return render_template('index.html')

@app.route('/register')
def register():
    return render_template('register.html')

@app.route('/review-administration')
@login_required
def review_administration():
    return render_template('administration.html')

@app.route('/review-management')
@login_required
def review_management():
    return render_template('management.html')

@app.route('/lecturer-dashboard')
@login_required
def lecturer_dashboard():
    return render_template('lecturerDashboard.html')

@app.route('/review-module')
@login_required
def review_module():
    return render_template('reviewModule.html')

@app.route('/owner')
@login_required
def owner():
    return render_template('owner.html')


@app.route('/login', methods=['POST'])
def login():
    email = request.form.get('email')
    password = request.form.get('password')
    
    getUser = text("SELECT UserID, Type, Email, Password FROM User WHERE Email = :Email")
    result = db.session.execute(getUser, {'Email': email})
    user = result.fetchone()

    if user:
        hashed_password_from_db = user.Password.encode('utf-8')

        # Check if the password matches the hashed password from the database
        if bcrypt.checkpw(password.encode('utf-8'), hashed_password_from_db):
            # Password is correct, set the user's session
            session['userID'] = user.UserID

            if user.Type == "Admin": 
                return redirect(url_for('review_administration'))
            elif user.Type == "Owner":
                return redirect(url_for('owner'))
            else:
                return redirect(url_for('lecturer_dashboard'))
        else:
            # Password is incorrect
            return redirect(url_for('home'))
    else:
        # User with the provided email doesn't exist
        return redirect(url_for('home'))


# REGISTER
@app.route('/register-user', methods=["POST"])
def register_user():
    try:
        name = request.form.get('name') 
        email = request.form.get('email')
        password = request.form.get('password')

        existing_user = db.session.execute(text("""SELECT email 
                                FROM User 
                                WHERE email = :email
                                """), {"email": email}).fetchone()
        
        password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())

        if (existing_user is None):
            # Insert user
            db.session.execute(text("""INSERT INTO User (Name, Email, Password, Type)
                                VALUES (:Name, :Email, :Password, :Type)"""), {"Name": name, "Email": email, "Password": password, "Type": "Lecturer"})
            db.session.commit()
            return redirect(url_for('home'))
        else: 
            return jsonify({'result': 'failure', 'error': 'Email already exists'}), 409  # HTTP 409 Conflict

    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# LOG OUT
@app.route('/logout')
def logout():
    session.pop('userID', None)
    return redirect(url_for('home'))


# DELEGATE REVIEWS
@app.route('/delegate-reviews', methods=['POST'])
def delegate_reviews():
    try:
        data = request.get_json()
        Completed = "To-Complete"

        for element in data['moduleID']:
            result = db.session.execute(text("""SELECT UserID 
                                            From Module 
                                            WHERE ModuleID = :ModuleID
                                            """), {"ModuleID": element})
            UserID = result.scalar()

            # Insert into Feedback table
            feedback_result = db.session.execute(text("""INSERT INTO Feedback (UserID, Deadline, Completed)
                                                    VALUES (:UserID, :Deadline, :Completed)"""), {"UserID": UserID, "Deadline": data['deadline'], "Completed": Completed})
            
            # Get FeedbackID from previous query
            FeedbackID = db.session.execute(text("SELECT LAST_INSERT_ID()")).fetchone()[0]

            # Insert into ModuleFeedback table
            db.session.execute(text("""INSERT INTO ModuleFeedback (FeedbackID, ModuleID)
                                        VALUES (:FeedbackID, :ModuleID)"""), {"FeedbackID": FeedbackID, "ModuleID": element})
            
            db.session.commit()

        return jsonify({'result': 'success'}), 200

    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# When module review is submitted
@app.route('/submit-review-endpoint', methods=['POST'])
def submit_review():
    data = request.get_json()
    UserID = 1

    try:
        FeedbackID = data['FeedbackID']
        
        # Update feedback record
        db.session.execute(text("""UPDATE feedback 
                                   SET UserID = :UserID, AcademicYear = :AcademicYear, School = :School, 
                                    Other = :Other, Date = :Date, Completed = :Completed, Author = :Author
                                   WHERE FeedbackID = :FeedbackID"""), {"UserID": UserID, "AcademicYear": data['academicYear'], "School": data['school'],"Other": data['other'], "Date": data['date'], "Completed": "Completed", "Author": data['author'], "FeedbackID": FeedbackID})

        # Update ModuleFeedback records associated with the given FeedbackID
        db.session.execute(text("""UPDATE ModuleFeedback
                                   SET StudentInfo = :StudentInfo, 
                                    ModuleEval = :ModuleEval, InclusiveNature = :InclusiveNature, 
                                    PastChanges = :PastChanges, FutureChanges = :FutureChanges, TeachingEval = :TeachingEval
                                   WHERE FeedbackID = :FeedbackID"""), {"FeedbackID": FeedbackID, "StudentInfo": data['studentInfo'], "ModuleEval": data['moduleEval'], "InclusiveNature": data['inclusiveNature'], "PastChanges": data['pastChanges'], "FutureChanges": data['futureChanges'], "TeachingEval": data['teachingEval']})

        db.session.commit()
        return jsonify({'result': 'success'}), 200
    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# When module review is saved
@app.route('/save-review-endpoint', methods=['POST'])
def save_review():
    data = request.get_json()
    UserID = session.get('userID')

    # As date cannot be empty
    if data['date'] == "":
        data['date'] = "2000-01-01"

    try:
        FeedbackID = data['FeedbackID']
        
        # Update feedback record
        db.session.execute(text("""UPDATE feedback 
                                   SET UserID = :UserID, AcademicYear = :AcademicYear, School = :School, 
                                    Other = :Other, Date = :Date, Completed = :Completed, Author = :Author
                                   WHERE FeedbackID = :FeedbackID"""), {"UserID": UserID, "AcademicYear": data['academicYear'], "School": data['school'],"Other": data['other'], "Date": data['date'], "Completed": "In-Progress", "Author": data['author'], "FeedbackID": FeedbackID})

        # Update ModuleFeedback records associated with the given FeedbackID
        db.session.execute(text("""UPDATE ModuleFeedback
                                   SET StudentInfo = :StudentInfo, 
                                    ModuleEval = :ModuleEval, InclusiveNature = :InclusiveNature, 
                                    PastChanges = :PastChanges, FutureChanges = :FutureChanges
                                   WHERE FeedbackID = :FeedbackID"""), {"FeedbackID": FeedbackID, "StudentInfo": data['studentInfo'], "ModuleEval": data['moduleEval'], "InclusiveNature": data['inclusiveNature'], "PastChanges": data['pastChanges'], "FutureChanges": data['futureChanges']})

        db.session.commit()
        return jsonify({'result': 'success'}), 200
    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# CALLED WHEN USER WANTS TO OPEN AN IN-PROGRESS REVIEW
@app.route('/edit-review')
def edit_review():
    feedbackID = request.args.get('FeedbackID')
    if not feedbackID:
        return "FeedbackID is required", 400

    try:
        feedbackID = int(feedbackID) 
        review_query = text("""SELECT f.*, mf.*
                               FROM feedback f
                               JOIN ModuleFeedback mf ON f.FeedbackID = mf.FeedbackID
                               WHERE f.FeedbackID = :FeedbackID""")
        
        review = db.session.execute(review_query, {'FeedbackID': feedbackID}).fetchone()

        if review:
            return render_template('editReview.html', review=review)
        else:
            return "Review not found", 404
    except ValueError:
        return "Invalid FeedbackID", 400
    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# GET USER REVIEW - USED TO DISPLAY TABLES WITH REVIEWS
@app.route('/get-user-reviews')
def get_reviews():
    UserID = session.get('userID')

    try:
        # Join Feedback with ModuleFeedback, then ModuleFeedback with Module
        all_reviews = text("""SELECT mo.ModuleName, f.Deadline, f.Completed, f.FeedbackID
                        FROM Feedback f
                        JOIN ModuleFeedback mf ON f.FeedbackID = mf.FeedbackID
                        JOIN Module mo ON mf.ModuleID = mo.ModuleID
                        WHERE f.UserID = :UserID""")
        result = db.session.execute(all_reviews, {'UserID': UserID}).mappings().all()

        all_reviews = [{'moduleName': row['ModuleName'], 'deadline': row['Deadline'], 'completed': row['Completed'], 'feedbackID': row['FeedbackID']} for row in result]

        return jsonify(all_reviews)

    except Exception as e:
        print(e)
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# GET ALL REVIEWS FOR ADMIN DASHBOARD
@app.route('/get-all-reviews')
def get_all_reviews():
    try:
        # Join Feedback with ModuleFeedback, then ModuleFeedback with Module
        all_reviews = text("""SELECT mo.ModuleName, mo.ModuleLead, f.Deadline, f.Completed, f.FeedbackID
                        FROM Feedback f
                        JOIN ModuleFeedback mf ON f.FeedbackID = mf.FeedbackID
                        JOIN Module mo ON mf.ModuleID = mo.ModuleID""")
        result = db.session.execute(all_reviews).mappings().all()

        all_reviews = [{'moduleName': row['ModuleName'], 'deadline': row['Deadline'], 'completed': row['Completed'], 'moduleLead': row['ModuleLead'], 'feedbackID': row['FeedbackID']} for row in result]

        return jsonify(all_reviews)

    except Exception as e:
        print(e)
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# GET MODULES FOR REVIEW DELEGATION
@app.route('/get-modules')
def get_modules():
    all_modules = text("""SELECT ModuleID, UserID, ModuleName, ModuleLead, ModuleCode, Credits
                       FROM Module""")

    result = db.session.execute(all_modules).mappings().all()

    all_modules = [{'moduleID': row['ModuleID'], 'UserID': row['UserID'], 'moduleName': row['ModuleName'], 'moduleLead': row['ModuleLead'], 'moduleCode': row['ModuleCode'], 'credits': row['Credits']} for row in result]

    return jsonify(all_modules)


# ADD MODULE INTO DATABASE
@app.route('/add-module', methods=['POST'])
def add_module():
    data = request.get_json()

    try:
        # FIND THE USERID OF THE MODULES LEAD LECTURER
        result = db.session.execute(text("""SELECT UserID
                                        FROM User
                                        WHERE Name = :ModuleLead"""), {'ModuleLead': data['moduleLead']})
        UserID_result = result.fetchone()

        if result is None:
            return jsonify({'result': 'failure', 'error': 'Module lead not found'}), 404

        UserID = UserID_result[0]

        add_result = db.session.execute(text("""INSERT INTO Module (ModuleCode, ModuleName, ModuleLead, Credits, UserID)
                                        VALUES (:ModuleCode, :ModuleName, :ModuleLead, :Credits, :UserID)"""), {"ModuleCode": data['moduleCode'], "ModuleName": data['moduleName'], "ModuleLead": data['moduleLead'], "Credits": data['credits'], "UserID": UserID})

        db.session.commit()

        return jsonify({'result': 'success'}), 200

    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500
    

# EDIT MODULES IN DATABASE
@app.route('/update-module-row', methods=['POST'])
def update_module_row():
    data = request.get_json()

    try:
        # FIND THE USERID OF THE MODULES LEAD LECTURER
        result = db.session.execute(text("""SELECT UserID
                                        FROM User
                                        WHERE Name = :ModuleLead"""), {'ModuleLead': data['moduleLead']})
        UserID_result = result.fetchone()

        if UserID_result is None:
            return jsonify({'result': 'failure', 'error': 'Module lead not found'}), 404

        UserID = UserID_result[0]

        add_result = db.session.execute(text("""UPDATE Module
                                        SET ModuleCode = :ModuleCode, ModuleName = :ModuleName, ModuleLead = :ModuleLead, Credits = :Credits, UserID = :UserID
                                        WHERE ModuleID = :ModuleID"""), {"ModuleCode": data['moduleCode'], "ModuleName": data['moduleName'], "ModuleLead": data['moduleLead'], "Credits": data['credits'], "UserID": UserID, "ModuleID": data['moduleID']})

        db.session.commit()

        return jsonify({'result': 'success'}), 200

    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# DELETE MODULE ROW
@app.route('/delete-module-row', methods=['POST'])
def delete_module_row():
    data = request.get_json()

    try:
        confilct = db.session.execute(text("""DELETE FROM ModuleFeedback
                            WHERE ModuleID = :ModuleID
                            """), {'ModuleID': data['moduleID']})

        conflictMF = db.session.execute(text("""DELETE FROM feedback
                            WHERE FeedbackID IN (
                                SELECT FeedbackID
                                FROM ModuleFeedback
                                WHERE ModuleID = :ModuleID
                            )
                            """), {'ModuleID': data['moduleID']})


        result = db.session.execute(text("""DELETE FROM Module
                                            WHERE ModuleID = :ModuleID"""), {"ModuleID": data['moduleID']})

        db.session.commit()

        return jsonify({'result': 'success'}), 200

    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500


# GET USERS FOR EMAIL DROPDOWN
@app.route('/get-users') 
def get_users():
    all_users = text("""SELECT UserID, Email, Name, Type
                       FROM User""")

    result = db.session.execute(all_users).mappings().all()

    all_users = [{'userID': row['UserID'], 'email': row['Email'], 'name': row['Name'], 'type': row['Type']} for row in result]

    return jsonify(all_users)


# EMAIL ALL USERS SELECTED
@app.route('/email-staff', methods=['POST'])
def email_staff():
    data = request.get_json()
    
    msg = Message(data['subject'],
                  sender="annualreviewsystem@gmail.com",
                  recipients=data['emails'],
                  body=data['message'])

    mail.send(msg)

    return jsonify({'result': 'success'}), 200


# EXPORT ALL SELECTED REVIEWS
@app.route('/export-word', methods=['POST'])
def export_word():
    from sqlalchemy import text
    data = request.get_json()
    
    FeedbackIDs = data['feedbackID']  # First element
    params = {"FeedbackID": tuple(FeedbackIDs)} 

    word_data = db.session.execute(text("""SELECT f.*, mf.*, m.*
                                    FROM feedback f
                                    JOIN ModuleFeedback mf ON f.FeedbackID = mf.FeedbackID
                                    JOIN Module m ON mf.ModuleID = m.ModuleID
                                    WHERE f.FeedbackID IN :FeedbackID OR mf.FeedbackID IN :FeedbackID
                                    """),
                                    params).mappings().fetchall()

    # Create Document
    document = Document()

    title_style = document.styles['Title']
    title_style.font.size = Pt(16)

    # HEADER should be - 61, 88, 151
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture('Images/Dundee_Logo-Word_Doc.jpg')

    column_widths = [Inches(2.5), Inches(4.5)]

    # FOOTER
    footer = document.sections[0].footer
    paragraph = footer.add_paragraph()

    # Set paragraph alignment and text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = "Quality and Academic Standards Office, September 2016"

    for run in paragraph.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x36, 0x5F, 0x91)

    paragraph2 = footer.add_paragraph("Updated February 2023")
    paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for run in paragraph2.runs:
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0x36, 0x5F, 0x91)

    # MAIN
    for row in word_data:
        title = document.add_paragraph('Annual Module Quality Enhancement Report', style='Title')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Display multiple fields in module details cell
        module_details = "Module Code: " + row.get("ModuleCode", "None") + \
                     "\nModule Name: " + row.get("ModuleName", "None") + \
                     "\nCredits: " + str(row.get("Credits", "None"))
        
        # Display date and author

        author_and_date = "Author: " + row.get("Author", "None") + \
                     "\nDate: " + str(row.get("Date", "None"))

        # Define the titles and variables
        records = {
            "1. Module details": module_details,
            "2. Academic Year": row.get("AcademicYear", "None"),
            "3. School": row.get("School", "None"),
            "4. Module Leader/Organiser": row.get("ModuleLead", "None"),
            "5. Student numbers, achievement and progression": row.get("StudentInfo", "None"),
            "6. Evaluation of the operation of the module": row.get("ModuleEval", "None"),
            "7. Evaluation of approach to teaching, assessment and feedback": row.get("TeachingEval", "None"),
            "8. Inclusive nature of the curriculum": row.get("InclusiveNature", "None"),
            "9. Effect of past changes": row.get("PastChanges", "None"),
            "10. Proposed future changes": row.get("FutureChanges", "None"),
            "11. Other comments": row.get("Other", "None"),
            "12. Author and date": author_and_date
        }

        table = document.add_table(rows=len(records), cols=2)
        table.style = 'Table Grid'

        # Set widths
        for row in table.rows:
            row.cells[0].width = column_widths[0]
            row.cells[1].width = column_widths[1]

        # Fill in the table 
        for i, (heading, text) in enumerate(records.items()):
            row_cells = table.rows[i].cells
            row_cells[0].text = heading
            row_cells[1].text = text

        document.add_page_break()


    document.save("Annual-Module-Quality-Enhancement-Reports.docx")

    # Return success response
    return {'result': 'success'}


# Change user to admin or lecturer
@app.route('/change-user-permissions', methods=['POST'])
def user_to_admin():
    data = request.get_json()

    try:
        update_user = db.session.execute(text("""UPDATE User
                                        SET Type = :Type
                                        WHERE UserID = :UserID"""), {"UserID": data['userID'], "Type": data['type']})

        db.session.commit()

        return jsonify({'result': 'success'}), 200

    except Exception as e:
        return jsonify({'result': 'failure', 'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True)